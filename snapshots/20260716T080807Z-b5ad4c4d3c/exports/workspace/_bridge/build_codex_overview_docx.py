#!/usr/bin/env python3
from __future__ import annotations

import math
import re
import shutil
import subprocess
import sys
from pathlib import Path

from PIL import Image, ImageChops
from docx import Document as DocxDocument


MERMAID_BLOCK_RE = re.compile(r"```mermaid\s*\n([\s\S]*?)```", re.M)


def image_width_in(path: Path, max_width: float = 6.2, max_height: float = 5.8) -> float:
    with Image.open(path) as img:
        w, h = img.size
    if w <= 0 or h <= 0:
        return max_width
    width_for_height = max_height * (w / h)
    return max(1.5, min(max_width, width_for_height))


def validate_rendered_images(images: list[Path]) -> None:
    for image_path in images:
        with Image.open(image_path) as image:
            image.load()
            width, height = image.size
            if width < 200 or height < 60:
                raise RuntimeError(
                    f"suspicious diagram dimensions: {image_path.name}={width}x{height}"
                )
            rgb = image.convert("RGB")
            white = Image.new("RGB", rgb.size, "white")
            content_bounds = ImageChops.difference(rgb, white).getbbox()
            if content_bounds is None:
                raise RuntimeError(f"blank rendered diagram: {image_path.name}")


def build_markdown(src_md: Path, image_dir: Path, out_md: Path) -> list[Path]:
    text = src_md.read_text(encoding="utf-8")
    images = sorted(p for p in image_dir.glob("*.png") if p.is_file())
    blocks = list(MERMAID_BLOCK_RE.finditer(text))
    if len(blocks) != len(images):
        raise RuntimeError(f"mermaid/image mismatch: {len(blocks)} blocks vs {len(images)} pngs")
    validate_rendered_images(images)

    chunks: list[str] = []
    pos = 0
    for idx, (block, img) in enumerate(zip(blocks, images), start=1):
        chunks.append(text[pos : block.start()])
        width = image_width_in(img)
        chunks.append(f"\n\n![Figure {idx}: {img.stem}]({img.name}){{width={width:.2f}in}}\n\n")
        pos = block.end()
    chunks.append(text[pos:])
    out_md.write_text("".join(chunks), encoding="utf-8")
    return images


def run_pandoc(input_md: Path, output_docx: Path) -> None:
    subprocess.run(
        [
            "pandoc",
            str(input_md),
            "--resource-path",
            str(input_md.parent),
            "-o",
            str(output_docx),
        ],
        check=True,
        cwd=input_md.parent,
    )


def enforce_image_caps(docx_path: Path, max_width_in: float = 6.2) -> None:
    doc = DocxDocument(docx_path)
    max_width_emu = int(max_width_in * 914400)
    changed = False
    for shape in doc.inline_shapes:
        if shape.width > max_width_emu:
            ratio = shape.height / shape.width
            shape.width = max_width_emu
            shape.height = int(max_width_emu * ratio)
            changed = True
    if changed:
        doc.save(docx_path)


def main() -> int:
    if len(sys.argv) != 4:
        print("Usage: build_codex_overview_docx.py <source.md> <render-dir> <output.docx>", file=sys.stderr)
        return 1

    src_md = Path(sys.argv[1]).resolve()
    render_dir = Path(sys.argv[2]).resolve()
    output_docx = Path(sys.argv[3]).resolve()
    if not src_md.is_file():
        print(f"missing source markdown: {src_md}", file=sys.stderr)
        return 2
    if not render_dir.is_dir():
        print(f"missing render dir: {render_dir}", file=sys.stderr)
        return 3

    temp_md = render_dir / "_codex-overview-pandoc.md"
    build_markdown(src_md, render_dir, temp_md)
    run_pandoc(temp_md, output_docx)
    enforce_image_caps(output_docx)

    doc = DocxDocument(output_docx)
    print(f"output={output_docx}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)}")
    for i, shape in enumerate(doc.inline_shapes, start=1):
        print(f"shape{i}={shape.width/914400:.2f}x{shape.height/914400:.2f}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
